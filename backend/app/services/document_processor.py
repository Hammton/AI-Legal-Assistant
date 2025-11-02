"""
Document Processing Service
Handles PDF, DOCX, and image extraction
"""
import logging
from typing import Optional
from pathlib import Path
import PyPDF2
import pdfplumber
from docx import Document

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Service for extracting text from various document formats"""
    
    async def extract_text(self, file_path: str) -> str:
        """
        Extract text from document based on file extension
        
        Args:
            file_path: Path to document file
            
        Returns:
            Extracted text content
        """
        path = Path(file_path)
        extension = path.suffix.lower()
        
        if extension == '.pdf':
            return await self.extract_pdf(file_path)
        elif extension in ['.docx', '.doc']:
            return await self.extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {extension}")
    
    async def extract_pdf(self, file_path: str) -> str:
        """
        Extract text from PDF using pdfplumber (fallback to PyPDF2)
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text
        """
        try:
            # Try pdfplumber first (better for complex PDFs)
            with pdfplumber.open(file_path) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n\n"
                
                if text.strip():
                    logger.info(f"Extracted {len(text)} characters from PDF using pdfplumber")
                    return text
        except Exception as e:
            logger.warning(f"pdfplumber failed: {e}, trying PyPDF2")
        
        try:
            # Fallback to PyPDF2
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n\n"
                
                logger.info(f"Extracted {len(text)} characters from PDF using PyPDF2")
                return text
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise
    
    async def extract_docx(self, file_path: str) -> str:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text
        """
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            logger.info(f"Extracted {len(text)} characters from DOCX")
            return text
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise
    
    async def extract_metadata(self, file_path: str) -> dict:
        """
        Extract document metadata
        
        Args:
            file_path: Path to document
            
        Returns:
            Metadata dictionary
        """
        path = Path(file_path)
        metadata = {
            "filename": path.name,
            "extension": path.suffix,
            "size_bytes": path.stat().st_size if path.exists() else 0
        }
        
        # Try to extract PDF metadata
        if path.suffix.lower() == '.pdf':
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    if pdf_reader.metadata:
                        metadata.update({
                            "title": pdf_reader.metadata.get('/Title', ''),
                            "author": pdf_reader.metadata.get('/Author', ''),
                            "subject": pdf_reader.metadata.get('/Subject', ''),
                            "creator": pdf_reader.metadata.get('/Creator', ''),
                            "pages": len(pdf_reader.pages)
                        })
            except Exception as e:
                logger.warning(f"Failed to extract PDF metadata: {e}")
        
        return metadata


# Global instance
document_processor = DocumentProcessor()
